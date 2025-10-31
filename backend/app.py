import os
import re
import json
import base64
import traceback
import pandas as pd

from pathlib import Path
from flask_cors import CORS
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from flask import Flask, request, jsonify, send_file, render_template, redirect, url_for, flash, session, abort, make_response, Response, send_from_directory

from utils.geo import nearest_on_route, reverse_geocode_neighborhood, haversine_m
from utils.gpx import parse_gpx, route_dists_grades, geojson_from_points
from utils.hashing import route_hash
from utils.physics import (
    infer_power_from_flat, speed_from_power,
    summarize_by_buckets
)
from utils.places import get_places_cached
from utils.map_video import generate_route_video

load_dotenv()

'''
Directories
'''

app = Flask(__name__, static_folder="static", static_url_path="/static")
'''
BASE = Path(__file__).resolve().parent
UPLOADS = BASE / "uploads"
UPLOADS.mkdir(exist_ok=True, parents=True)
CACHE_DIR = BASE / "cache" / "places"
CACHE_DIR.mkdir(exist_ok=True, parents=True)
REPORT_DIR = BASE / "static" / "reports"
REPORT_DIR.mkdir(parents=True, exist_ok=True)
'''

BASE = Path(__file__).resolve().parent
UPLOADS = BASE / "uploads"
CACHE_DIR = BASE / "cache" / "places"
CACHE_DIR.mkdir(exist_ok=True, parents=True)
REPORT_DIR = BASE / "static" / "reports"
REPORT_DIR.mkdir(parents=True, exist_ok=True)
STATIC_VID = BASE / "static" / "videos"
STATIC_VID.mkdir(parents=True, exist_ok=True)

#app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB GPX
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev")

# CORS (React dev server)
origins = os.environ.get("CORS_ORIGINS", "http://localhost:5173").split(",")
CORS(app, resources={
    r"/api/*": {"origins": "*"},
    r"/static/videos/*": {"origins": "*"},
})

_data_url_re = re.compile(r"^data:image/(png|jpeg);base64,(.*)$", re.I)

"""
Function to save the route (helps save time when reloading the same route)
"""
def _save_route(file_id, payload):
    (UPLOADS / f"{file_id}.route.json").write_text(json.dumps(payload))

"""
Loads the route once a new one, or the same one is requested
"""
def _load_route(file_id):
    p = UPLOADS / f"{file_id}.route.json"
    if not p.exists():
        return None
    return json.loads(p.read_text())

def _load_cached(file_id: str):
    j1 = CACHE_DIR / f"{file_id}.json"
    if j1.exists():
        return json.loads(j1.read_text())

    # 🔧 try route hash or "_result" fallback
    for suffix in ["_result.json", "_analysis.json"]:
        j2 = CACHE_DIR / f"{file_id}{suffix}"
        if j2.exists():
            return json.loads(j2.read_text())

    # also look for files starting with this prefix
    matches = list(CACHE_DIR.glob(f"{file_id}*.json"))
    if matches:
        return json.loads(matches[0].read_text())

    raise FileNotFoundError(f"Cached analysis not found for {file_id}")


"""
Analysis/Predictions are made with regards to the route
"""
def _compute_analysis(points, params):
    dists, grades, egain, eloss = route_dists_grades(points)
    total_m = sum(dists)

    flat_mph = float(params.get("flat_mph", 15.0))
    base_mps = flat_mph * 0.44704

    physics = bool(params.get("physics", False))
    rider_kg = float(params.get("rider_kg", 75.0))
    bike_kg  = float(params.get("bike_kg",  10.0))
    cda      = float(params.get("cda",       0.30))
    crr      = float(params.get("crr",       0.004))
    wind_mps = float(params.get("wind_mph",  0.0)) * 0.44704
    m_total  = rider_kg + bike_kg

    speeds = []
    if physics:
        # infer constant power from flat target pace
        P = infer_power_from_flat(base_mps, m_total, cda, crr)
        for g in grades:
            v = speed_from_power(P, g, m_total, cda, crr, wind_mps=wind_mps)
            speeds.append(v)
    else:
        # simple heuristic model
        alpha_up, alpha_down = 12.0, 2.0
        min_v, max_v = 3.0*0.44704, 35.0*0.44704
        for g in grades:
            if g >= 0:
                v = base_mps * max(0.2, (1 - alpha_up * g))
            else:
                v = base_mps * (1 + alpha_down * (-g))
            speeds.append(max(min_v, min(max_v, v)))

    times = [d / max(v, 0.1) for d, v in zip(dists, speeds)]
    total_s = sum(times)
    avg_mph = (total_m/1609.34) / (total_s/3600.0) if total_s > 0 else 0.0

    # charts downsample
    cum = []
    s = 0.0
    for d in dists:
        s += d
        cum.append(s)
    step = max(1, len(dists)//1500)
    chart = {
        "dist_mi": [round(x/1609.34, 2) for x in cum[::step]],
        "speed_mph": [round(v*2.23694, 2) for v in speeds[::step]],
        "grade_pct": [round(g*100.0, 2) for g in grades[::step]],
    }

    # splits by grade buckets
    splits = summarize_by_buckets(dists, grades, speeds)

    return {
        "distance_mi": round(total_m/1609.34, 2),
        "distance_km": round(total_m/1000.0, 2),
        "elev_gain_ft": round(egain*3.28084, 0),
        "elev_gain_m": round(egain, 1),
        "pred_time_h": round(total_s/3600.0, 2),
        "pred_avg_mph": round(avg_mph, 2),
        "flat_target_mph": round(flat_mph, 2),
        "physics_used": physics,
        "chart": chart,
        "splits": splits,
        "params_used": {
            "rider_kg": rider_kg, "bike_kg": bike_kg, "cda": cda, "crr": crr,
            "wind_mph": params.get("wind_mph", 0.0)
        }
    }

"""
Helper function to save the generated plots as images when the user requests to download the report
"""
def _save_data_url_png(data_url: str, dest: Path):
    """Save a dataURL 'data:image/png;base64,...' to a file."""
    m = _data_url_re.match(data_url or "")
    if not m:
        raise ValueError("Invalid image data URL")
    b64 = m.group(2)
    dest.parent.mkdir(parents=True, exist_ok=True)
    with open(dest, "wb") as f:
        f.write(base64.b64decode(b64))

"""
Helper function for empty/null values
"""
def _fnum(x):
    if x is None:
        return None
    try:
        return float(str(x).replace(",", "").strip())
    except Exception:
        return None

"""
API call to perform the analysis of the route before it is even done
"""
@app.post("/api/analyze")
def api_analyze():
    if "gpx" not in request.files:
        return jsonify({"error": "No GPX uploaded"}), 400
    f = request.files["gpx"]
    file_id = secure_filename(f.filename or "route.gpx")
    path = UPLOADS / file_id
    f.save(path)

    points = parse_gpx(path)
    if len(points) < 2:
        return jsonify({"error": "Not enough track points"}), 400

    # compute hash + store raw
    rhash = route_hash(points)
    route_payload = {"file_id": file_id, "route_hash": rhash, "points": points}
    _save_route(file_id, route_payload)

    # params
    params = request.form.to_dict()
    analysis = _compute_analysis(points, params)

    # geojson downsampled for map
    gjson = geojson_from_points(points, step=max(1, len(points)//1500))

    return jsonify({
        "file_id": file_id,
        "route_hash": rhash,
        "geojson": gjson,
        "analysis": analysis
    })

def _read_json(p: Path):
    try:
        return json.loads(p.read_text())
    except Exception:
        return None

# your existing _load_cached(...) can stay as-is if you want the fallback

@app.post("/api/video")
def api_video():
    body = request.get_json(force=True)

    # 1) Prefer direct payload (no-cache path)
    geojson = body.get("geojson")
    dist_mi = body.get("dist_mi")
    speed_mph = body.get("speed_mph")
    elev_ft = body.get("elev_ft")
    use_tiles = bool(body.get("map_tiles", False))
    padding = float(body.get("padding", 0.05))

    # 2) Fallback to cache if not provided
    if not (geojson and dist_mi and speed_mph):
        file_id = body.get("file_id") or body.get("route_hash") or body.get("route_name")
        if not file_id:
            return jsonify({"error": "Provide geojson/dist_mi/speed_mph, or a file_id/route_hash"}), 400

        try:
            cached = _load_cached(file_id)  # your tolerant loader
        except Exception as e:
            return jsonify({"error": f"Load failed: {e}"}), 400

        geojson = cached.get("geojson")
        analysis = (cached or {}).get("analysis", {})
        chart = (analysis or {}).get("chart", {})
        dist_mi = chart.get("dist_mi")
        speed_mph = chart.get("speed_mph")

    # Validate arrays
    if not geojson or not dist_mi or not speed_mph or len(dist_mi) != len(speed_mph):
        return jsonify({"error": "Missing or inconsistent chart arrays (dist_mi/speed_mph)"}), 400

    out_name = f"{(body.get('route_name') or body.get('route_hash') or 'route')}.mp4"
    out_path = STATIC_VID / out_name

    try:
        generate_route_video(
            geojson=geojson,
            dist_mi=dist_mi,
            speed_mph=speed_mph,
            out_path=str(out_path),
            fps=int(body.get("fps", 30)),
            width=int(body.get("width", 1280)),
            height=int(body.get("height", 720)),
            max_seconds=int(body.get("max_seconds", 60)),
            map_tiles=use_tiles,         
            padding=padding,  
        )
    except Exception as e:
        return jsonify({"error": f"Video generation error: {e}"}), 500

    return jsonify({"url": f"/static/videos/{out_name}"})

@app.get("/api/video/download/<filename>")
def api_video_download(filename: str):
    """Streams MP4 with Content-Disposition: attachment (prompts download)."""
    mp4 = STATIC_VID / filename
    if not mp4.exists():
        return jsonify({"error": "File not found"}), 404
    # If you prefer, use send_from_directory(..., as_attachment=True)
    return send_file(
        mp4,
        mimetype="video/mp4",
        as_attachment=True,
        download_name=filename,
        conditional=True,  # supports range requests/resume
    )

"""
Recomputes/refreshes the results page if the user makes some changes of values to the form
"""
@app.post("/api/recompute")
def api_recompute():
    data = request.get_json(force=True)
    file_id = data.get("file_id")
    route = _load_route(file_id)
    if not route:
        return jsonify({"error": "Unknown file_id"}), 404
    params = data.get("params", {})
    analysis = _compute_analysis(route["points"], params)
    return jsonify({"file_id": file_id, "route_hash": route["route_hash"], "analysis": analysis})

"""
Renders the result page, and outputs the respective charts and graphics
"""
@app.get("/api/result")
def api_result():
    file_id = request.args.get("file_id")
    route = _load_route(file_id)
    if not route:
        return jsonify({"error": "Unknown file_id"}), 404
    # default analysis if someone hits deep link with no params: 15 mph, no physics
    default_params = {"flat_mph": 15.0, "physics": False}
    analysis = _compute_analysis(route["points"], default_params)
    gjson = geojson_from_points(route["points"], step=max(1, len(route["points"])//1500))
    return jsonify({"file_id": file_id, "route_hash": route["route_hash"], "geojson": gjson, "analysis": analysis})

"""
API call to display the locations along the route within a 1 mile radius of the route
"""
@app.get("/api/places")
def api_places():
    file_id = request.args.get("file_id")
    route = _load_route(file_id)
    if not route:
        return jsonify({"error": "Unknown file_id"}), 404

    points = route["points"]
    # sample every ~50 points for cost control
    samples = [(lat, lon) for i, (lat, lon, _) in enumerate(points) if i % 50 == 0 or i == len(points)-1]
    api_key = os.environ.get("GOOGLE_MAPS_API_KEY", "")
    data = get_places_cached(CACHE_DIR, route["route_hash"], api_key, samples)
    return jsonify({"route_hash": route["route_hash"], **data})

"""
Downloads any file, the user wants from the site
"""
@app.get("/api/download")
def api_download():
    file_id = request.args.get("file_id")
    p = UPLOADS / file_id
    if not p.exists():
        return jsonify({"error": "Not found"}), 404
    return send_file(p, as_attachment=True, download_name=file_id)

"""
Displays the charts on the results page
"""
@app.post("/api/charts")
def api_save_charts():
    try:
        j = request.get_json(force=True)
        file_id = j.get("file_id")
        speed_png = j.get("speed_png")
        grade_png = j.get("grade_png")
        if not file_id or not speed_png or not grade_png:
            return jsonify({"error": "file_id, speed_png, grade_png required"}), 400

        out = REPORT_DIR / file_id
        _save_data_url_png(speed_png, out / "speed.png")
        _save_data_url_png(grade_png, out / "grade.png")
        return jsonify({"ok": True})
    except Exception as e:
        print("api_save_charts error:\n", traceback.format_exc())
        return jsonify({"error": f"charts: {type(e).__name__}: {e}"}), 500

"""
Generates a report if the user wishes summarizing everything on the route
"""
@app.get("/api/report")
def api_report():
    try:
        file_id = request.args.get("file_id")
        if not file_id:
            return jsonify({"error": "file_id required"}), 400

        # Load route (points: [(lat, lon, elev), ...])
        route = _load_route(file_id)
        if not route:
            return jsonify({"error": "Unknown file_id"}), 404

        # Default analysis (heuristic model; keep your defaults)
        analysis = _compute_analysis(
            route["points"], {"flat_mph": 15.0, "physics": False}
        )

        # Build cumulative distances along route (meters)
        dists, grades, *_ = route_dists_grades(route["points"])
        cum = [0.0]
        for d in dists:
            cum.append(cum[-1] + d)
        total_m = cum[-1]  # total route length (meters)

        # Places (cached by route hash or fetch-then-cache)
        samples = [
            (lat, lon)
            for i, (lat, lon, _) in enumerate(route["points"])
            if i % 50 == 0 or i == len(route["points"]) - 1
        ]
        api_key = os.environ.get("GOOGLE_MAPS_API_KEY", "")
        places_data = get_places_cached(CACHE_DIR, route["route_hash"], api_key, samples)
        pois = places_data.get("places", [])

        # Enrich each POI with:
        #  - off_route_miles: straight-line distance to nearest route point
        #  - miles_to_finish: route miles remaining from the nearest snapped location
        #  - neighborhood: reverse-geocoded borough/city (fallback to vicinity)
        geocache = CACHE_DIR / "reverse_geocode"
        geocache.mkdir(parents=True, exist_ok=True)

        enriched = []
        for p in pois:
            plat, plon = p.get("lat"), p.get("lon")
            off_m, along_m, _ = nearest_on_route(route["points"], cum, plat, plon)
            miles_off = off_m / 1609.34
            miles_to_end = max(0.0, (total_m - along_m)) / 1609.34

            # neighborhood via reverse geocode (cached). If no key, may return None.
            neigh = reverse_geocode_neighborhood(plat, plon, api_key, geocache)
            enriched.append(
                {
                    **p,
                    "off_route_miles": round(miles_off, 2),
                    "miles_to_finish": round(miles_to_end, 2),
                    "neighborhood": neigh or p.get("neighborhood") or p.get("vicinity"),
                }
            )

        # Optional video link if already generated
        video_path = UPLOADS / "videos" / f"{file_id}.mp4"
        video_url = f"/static/uploads/videos/{file_id}.mp4" if video_path.exists() else None

        # Build DOCX
        rep_dir = UPLOADS / "reports"
        rep_dir.mkdir(parents=True, exist_ok=True)
        out = rep_dir / f"{file_id}.docx"

        doc = Document()
        doc.add_heading("Route Plan", level=1)
        doc.add_paragraph(f"File: {file_id}")
        doc.add_paragraph(
            f"Distance: {analysis['distance_mi']} mi ({analysis['distance_km']} km)"
        )
        doc.add_paragraph(f"Elevation gain: {analysis['elev_gain_ft']} ft")
        doc.add_paragraph(f"Predicted time (heuristic): {analysis['pred_time_h']} h")
        doc.add_paragraph(
            f"Predicted average speed: {analysis['pred_avg_mph']} mph"
        )
        if video_url:
            doc.add_paragraph(f"Route video: {video_url}")

        # ---- Embed charts if uploaded ----
        img_dir = REPORT_DIR / file_id
        speed_png = img_dir / "speed.png"
        grade_png = img_dir / "grade.png"

        doc.add_heading("Speed & Grade", level=2)
        if speed_png.exists():
            doc.add_paragraph("Speed (mph) vs. Distance")
            doc.add_picture(str(speed_png), width=Inches(6.5))
        if grade_png.exists():
            doc.add_paragraph("Grade (%) vs. Distance")
            doc.add_picture(str(grade_png), width=Inches(6.5))

        # ---- Splits table ----
        doc.add_heading("Splits by Grade Bucket", level=2)
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Grade Bucket"
        hdr[1].text = "Miles"
        hdr[2].text = "Avg mph"
        hdr[3].text = "Pace (min/mi)"
        for r in analysis["splits"]:
            row = table.add_row().cells
            row[0].text = str(r["bucket"])
            row[1].text = str(r["miles"])
            row[2].text = str(r["avg_mph"])
            row[3].text = str(r.get("pace_min_per_mile") or "-")

        # ---- ALL places, grouped by category ----
        doc.add_heading("All Points of Interest", level=2)
        by_cat = {}
        for p in enriched:
            by_cat.setdefault(p.get("category", "other"), []).append(p)

        for cat, arr in by_cat.items():
            doc.add_heading(cat.replace("_", " ").title(), level=3)
            t = doc.add_table(rows=1, cols=5)
            h = t.rows[0].cells
            h[0].text = "Name"
            h[1].text = "Neighborhood"
            h[2].text = "Vicinity"
            h[3].text = "Off-route (mi)"
            h[4].text = "Miles to finish"
            for p in arr:
                r = t.add_row().cells
                r[0].text = p.get("name") or "-"
                r[1].text = (p.get("neighborhood") or "-")
                r[2].text = (p.get("vicinity") or "-")
                r[3].text = f"{p.get('off_route_miles', 0)}"
                r[4].text = f"{p.get('miles_to_finish', 0)}"

        doc.save(out)
        return send_file(out, as_attachment=True, download_name=out.name)
    except Exception as e:
        print("api_report error:\n", traceback.format_exc())
        return jsonify({"error": f"report: {type(e).__name__}: {e}"}), 500

"""
Generates an excel/csv file of all the locations to stop at
"""
@app.post("/api/stops_xlsx")
def api_stops_xlsx():
    try:
        j = request.get_json(force=True)
        file_id = j.get("file_id")
        places = j.get("places") or []
        if not file_id or not isinstance(places, list):
            return jsonify({"error": "file_id and places[] required"}), 400

        rows = []
        for p in places:
            # Normalize keys to match your report columns exactly
            category = p.get("category") or p.get("type") or ""
            name = p.get("name") or ""
            vicinity = p.get("vicinity") or p.get("address") or ""
            neighborhood = (
                p.get("neighborhood") or p.get("city") or p.get("locality") or ""
            )

            off_route = (
                p.get("off_route_mi")
                or p.get("off_miles")
                or p.get("off_mi")
                or p.get("distance_off_mi")
            )
            miles_finish = (
                p.get("miles_to_finish")
                or p.get("to_finish_mi")
                or p.get("dist_to_finish_mi")
                or p.get("mi_remaining")
            )

            rows.append({
                "Category":        str(category),
                "Name":            str(name),
                "Neighborhood":    str(neighborhood),
                "Vicinity":        str(vicinity),
                "Off-route (mi)":  _fnum(off_route),
                "Miles to finish": _fnum(miles_finish),
            })

        df = pd.DataFrame(rows)

        # Desired column order (Category optional)
        cols = ["Name", "Neighborhood", "Vicinity", "Off-route (mi)", "Miles to finish"]
        if df["Category"].astype(str).str.len().gt(0).any():
            cols = ["Category"] + cols
        df = df.reindex(columns=cols)

        # Numeric cleanup + sort from start -> finish
        for c in ("Off-route (mi)", "Miles to finish"):
            if c in df.columns:
                df[c] = pd.to_numeric(df[c], errors="coerce")
        if "Miles to finish" in df.columns:
            df = df.sort_values("Miles to finish", ascending=False, kind="stable")

        out_dir = REPORT_DIR / file_id
        out_dir.mkdir(parents=True, exist_ok=True)
        xlsx_path = out_dir / "stops.xlsx"
        csv_path  = out_dir / "stops.csv"

        df.to_excel(xlsx_path, index=False, engine="openpyxl")
        df.to_csv(csv_path, index=False)

        return jsonify({
            "xlsx": f"/static/reports/{file_id}/stops.xlsx",
            "csv":  f"/static/reports/{file_id}/stops.csv",
        })
    except Exception as e:
        print("stops_xlsx error:\n", traceback.format_exc())
        return jsonify({"error": f"stops_xlsx: {type(e).__name__}: {e}"}), 500

"""
main function to run the program
"""
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5001, debug=True)
